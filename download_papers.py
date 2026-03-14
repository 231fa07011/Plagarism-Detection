import os
import requests
import urllib3
from docx import Document

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

out_dir = r"c:\Users\MY PC\OneDrive\Desktop\Plagarism-Detection\sample_test_files\real_papers"
os.makedirs(out_dir, exist_ok=True)

downloads = [
    {
        "url": "https://raw.githubusercontent.com/jbrownlee/Datasets/master/review_polarity/txt_sentoken/pos/cv000_29590.txt",
        "filename": "A3_Article_Sample_Review.txt",
        "type": "TXT Article"
    },
    {
        "url": "https://raw.githubusercontent.com/jbrownlee/Datasets/master/pima-indians-diabetes.names",
        "filename": "A4_Pima_Indians_Diabetes.txt",
        "type": "TXT Article"
    },
    {
        "url": "https://raw.githubusercontent.com/jbrownlee/Datasets/master/housing.names",
        "filename": "A5_Boston_Housing.txt",
        "type": "TXT Article"
    }
]

print("Starting downloads from official academic datasets...")

for item in downloads:
    try:
        output_path = os.path.join(out_dir, item["filename"])
        print(f"Downloading {item['type']}: {item['filename']}...")
        
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(item["url"], headers=headers, stream=True, timeout=(10, 15), verify=False)
        response.raise_for_status()
        
        with open(output_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        print(" -> Success!")
    except Exception as e:
        print(f" -> Failed: {e}")

print("Generating Document: Sample_Official_Document_A3.docx...")
try:
    docx_path = os.path.join(out_dir, "Sample_Official_Document_A3.docx")
    doc = Document()
    doc.add_heading('Official Document Example: A3', 0)
    doc.add_paragraph('This is a locally generated official document representing A3 standard text.')
    doc.save(docx_path)
    print(" -> Success!")
except Exception as e:
    print(f" -> Failed: {e}")

print("Generating Document: Sample_Official_Document_A4.docx...")
try:
    docx_path = os.path.join(out_dir, "Sample_Official_Document_A4.docx")
    doc = Document()
    doc.add_heading('Official Document Example: A4', 0)
    doc.add_paragraph('This is a locally generated official document representing A4 standard text.')
    doc.save(docx_path)
    print(" -> Success!")
except Exception as e:
    print(f" -> Failed: {e}")

print("Generating Document: Sample_Official_Document_A5.docx...")
try:
    docx_path = os.path.join(out_dir, "Sample_Official_Document_A5.docx")
    doc = Document()
    doc.add_heading('Official Document Example: A5', 0)
    doc.add_paragraph('This is a locally generated official document representing A5 standard text.')
    doc.save(docx_path)
    print(" -> Success!")
except Exception as e:
    print(f" -> Failed: {e}")

print("\nAll downloads and generation completed successfully.")
