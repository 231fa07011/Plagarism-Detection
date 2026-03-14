"""
MASTER SAMPLE GENERATOR v2.0
Generates 7 sets of TXT, PDF, and DOCX files across different plagiarism scenarios.
"""
import os
import subprocess
import sys

# --- AUTO-INSTALL DEPENDENCIES ---
def setup_environment():
    pkgs = ['fpdf2', 'python-docx']
    for pkg in pkgs:
        try:
            if pkg == 'fpdf2': import fpdf
            elif pkg == 'python-docx': import docx
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "--quiet"])

setup_environment()

from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIG ---
# Get the directory where this script is located
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.join(ROOT_DIR, "sample_test_files")
SUBFOLDERS = ["exact_copy", "paraphrasing", "all_flagged", "original_content"]

for sub in SUBFOLDERS:
    os.makedirs(os.path.join(BASE_DIR, sub), exist_ok=True)

# Content Bank with deeply detailed academic sections
SCENARIOS = {
    "exact_copy": [
        ("AI_Evolution", (
            "Introduction: Artificial intelligence (AI) has emerged as a cornerstone of modern technological progress, specifically reconfiguring the educational landscape by enabling personalized learning paths tailored to individual student needs.\n\n"
            "Body 1: At its core, machine learning - a critical subset of AI - leverages sophisticated algorithms to parse vast datasets, identifying patterns that allow systems to improve their predictive accuracy autonomously without being explicitly programmed for every scenario.\n\n"
            "Body 2: Modern pedagogy is increasingly incorporating automated assessment tools that provide immediate, actionable feedback. These systems assist lecturers in identifying learning gaps early, allowing for targeted intervention and a more dynamic classroom environment.\n\n"
            "Data: Recent longitudinal studies demonstrate a 58% improvement in concept mastery when students utilize adaptive AI tutors, compared to a mere 12% improvement in standard lecture-based settings.\n\n"
            "Conclusion: Ultimately, the long-term success of academia hinges on the ethical and strategic integration of these intelligent systems to amplify human potential rather than replace it."
        )),
        ("Blockchain_Tech", (
            "Introduction: Blockchain technology represents a radical departure from centralized database management, offering a robust decentralized framework that ensures data integrity and operational transparency through distributed ledgers.\n\n"
            "Body 1: Each block in the sequence is linked via a cryptographic hash of the preceding block, creating a chronological chain that is virtually impossible to alter without the consensus of the entire network.\n\n"
            "Body 2: The implementation of smart contracts - self-executing agreements with terms written directly into code - further removes the need for traditional intermediaries like banks or legal firms, drastically reducing transaction costs.\n\n"
            "Data: Since 2020, investment in blockchain-based supply chain solutions has grown by 300% annually, with over $15 billion currently locked in decentralized liquidity pools.\n\n"
            "Conclusion: While scalability remains a challenge, the shift toward zero-trust distributed systems marks the most significant evolution in global record-keeping since the advent of digital computing."
        )),
        ("Neural_Nets", (
            "Introduction: Artificial neural networks (ANNs) are sophisticated computational models designed to mimic the structural and functional characteristics of biological neural pathways found in the human brain.\n\n"
            "Body 1: These systems consist of thousands of interconnected processing nodes arranged in layers. Each connection is weighted, and these weights are fine-tuned during the training process to minimize the deviation from expected results.\n\n"
            "Body 2: Advanced techniques such as Deep Learning utilize multiple hidden layers to extract high-level features from raw data, enabling capabilities like facial recognition, natural language processing, and medical diagnosis.\n\n"
            "Data: The most advanced Large Language Models (LLMs) currently operate on architectures featuring over 1.5 trillion individual parameters, requiring massive parallel processing power.\n\n"
            "Conclusion: As hardware specialized for tensor operations continues to improve, the gap between synthetic pattern matching and human cognitive flexibility is rapidly narrowing."
        )),
        ("Quantum_Limit", (
            "Introduction: Quantum computing is poised to redefine the limits of computational complexity by utilizing the principles of quantum mechanics - specifically superposition and entanglement - to process information in ways classical bits cannot.\n\n"
            "Body 1: Unlike classical systems that rely on binary states (0 or 1), a quantum bit (qubit) can exist in multiple states simultaneously, allowing for the concurrent exploration of vast solution spaces for complex problems.\n\n"
            "Body 2: Currently, the biggest hurdle is maintaining 'coherence' - the stable state of qubits. Researchers are developing ultra-cold environments near absolute zero to prevent external noise from disrupting the delicate quantum calculations.\n\n"
            "Data: In recent benchmarks, a 53-qubit quantum processor successfully completed a random circuit sampling task in approximately 200 seconds, a feat estimated to take the world's fastest supercomputer 10,000 years.\n\n"
            "Conclusion: Achieving broad quantum supremacy will revolutionize fields ranging from cryptography and drug discovery to fundamental materials science."
        )),
        ("IoT_Future", (
            "Introduction: The Internet of Things (IoT) encompasses a vast network of physical objects embedded with sensors and software for the purpose of connecting and exchanging data with other devices over the internet.\n\n"
            "Body 1: From smart thermostats to industrial turbines, these connected devices generate a continuous stream of telemetry data, enabling predictive maintenance and real-time optimization of complex infrastructure.\n\n"
            "Body 2: As the volume of data grows, 'Edge Computing' has become vital. Processing data locally at the device level reduces latency and bandwidth usage, ensuring that critical safety systems can react instantly without waiting for cloud processing.\n\n"
            "Data: Industry analysts forecast that by 2030, there will be more than 125 billion connected IoT devices worldwide, generating over 80 zettabytes of data annually.\n\n"
            "Conclusion: While the hyper-connected future promises unprecedented efficiency, it also necessitates a new paradigm for cybersecurity to protect critical infrastructure from remote exploitation."
        )),
        ("Cyber_Defense", (
            "Introduction: In an era of escalating digital threats, the Zero Trust security model has emerged as a fundamental architectural shift, operating on the principle of 'never trust, always verify' for every connection attempt.\n\n"
            "Body 1: Traditional perimeter-based security is no longer sufficient against sophisticated persistent threats. Modern defense requires granular access controls and continuous monitoring of all user and device behavior within the network.\n\n"
            "Body 2: Multi-factor authentication (MFA) and end-to-end encryption form the core pillars of a resilient defense strategy. However, human factors remain the weakest link, necessitating rigorous security awareness training for all personnel.\n\n"
            "Data: The average cost of a single data breach in the enterprise sector now exceeds $4.2 million, with 85% of successful attacks involving some form of social engineering or phishing.\n\n"
            "Conclusion: Effective cybersecurity is not a one-time product installation but a continuous process of adaptation, vigilance, and proactive threat hunting."
        )),
        ("Cloud_Systems", (
            "Introduction: Cloud computing has fundamentally transformed how businesses deploy and scale digital services, providing on-demand access to a shared pool of configurable computing resources with minimal management effort.\n\n"
            "Body 1: The major service models - Infrastructure as a Service (IaaS), Platform as a Service (PaaS), and Software as a Service (SaaS) - allow organizations to shift from heavy capital expenditure to flexible operational expenditure.\n\n"
            "Body 2: Cloud-native architectures, utilizing microservices and container orchestration (like Kubernetes), provide the agility needed for rapid deployment and automated scaling in response to fluctuating market demands.\n\n"
            "Data: Public cloud spending is expected to reach $600 billion by 2024, as nearly 94% of modern enterprises have migrated at least 60% of their workloads to the cloud.\n\n"
            "Conclusion: The migration to cloud-centric infrastructure is a prerequisite for digital transformation, enabling the rapid innovation required in the modern global economy."
        ))
    ],
    "paraphrasing": [
        ("AI_Rephrased", (
            "Introduction: Educational institutions are currently being fundamentally altered by AI-driven innovations that permit the creation of specialized learning experiences for students in various disciplines.\n\n"
            "Body 1: The core mechanism involves systems that can detect complex data relationships and refine their own accuracy over time through intensive observation, rather than following rigid manual instructions.\n\n"
            "Body 2: Higher education is increasingly adopting automated feedback loops that give students and teachers immediate insights into performance. This helps identify where extra support is needed much faster than traditional methods.\n\n"
            "Data: Information gathered from pilot programs suggests that using smart tutoring apps can increase student focus by roughly 40% compared to standard classroom environments.\n\n"
            "Conclusion: Ultimately, the evolution of contemporary schooling will be determined by how safely we can merge human teaching skills with automated technical capabilities."
        )),
        ("Blockchain_Rephrased", (
            "Introduction: Distributed ledger frameworks offer a complete break from traditional central authority models, providing a safe way to store records using a network of spread-out computers.\n\n"
            "Body 1: Every new entry in the sequence is secured by a unique code tied to the previous record, making the entire history permanent and resistant to fraudulent changes by any single party.\n\n"
            "Body 2: The use of programmable deal-making code - often referred to as smart contracts - removes the requirement for traditional middle-men like lawyers or bankers, which significantly lowers the cost of doing business.\n\n"
            "Data: Statistically, over $20 billion is now moving through decentralized financial networks every month, a massive increase from just a few years ago.\n\n"
            "Conclusion: Even with the current technical hurdles, the move toward distributed trust-building is the most important change in financial records since the invention of the computer."
        )),
        ("Brain_Computing", (
            "Introduction: Synthetic neural architectures are advanced mathematical systems that attempt to replicate the way biological brains process complex sensory information and learn from it.\n\n"
            "Body 1: These platforms comprise thousands of interconnected points of processing that work together in layers. Each link has a specific 'weight' that changes as the system tries to get closer to a correct answer.\n\n"
            "Body 2: Deep learning strategies use multiple layers to pull focus onto specific features of data, allowing machines to do things like see objects in photos or understand spoken human language with high precision.\n\n"
            "Data: Leading researchers have built language systems that use trillions of unique connections, requiring incredible amounts of energy and server space to maintain.\n\n"
            "Conclusion: As specialized computer chips for math-heavy tasks get more powerful, the line between machine calculation and human-like understanding is becoming very blurry."
        )),
        ("Quantum_Rephrased", (
            "Introduction: Computing with quantum mechanics is set to change what we thought was possible in mathematics by using the strange properties of subatomic particles to solve previously impossible riddles.\n\n"
            "Body 1: Instead of just using zeros and ones, these machines use bits that can be both at the same time. This allows them to search for thousands of solutions to a problem at once.\n\n"
            "Body 2: The biggest current obstacle is keeping the machine stable and cold enough to work. Even a tiny bit of heat or outside noise can cause the calculation to fail instantly.\n\n"
            "Data: In recent tests, a quantum processor finished a task in under four minutes that would have taken a traditional super-powered computer many thousands of years.\n\n"
            "Conclusion: Full mastery of these machines will change everything from how we protect our data to how we discover new life-saving medicines."
        )),
        ("Smart_Homes", (
            "Introduction: The network of connected household items allows for a seamless exchange of data between appliances over a home Wi-Fi network to make life more convenient.\n\n"
            "Body 1: Devices ranging from heaters to smart lights are constant sources of information that help homeowners track how much power they are using and automate their daily routines through apps.\n\n"
            "Body 2: To keep things running fast, many of these systems now do their 'thinking' inside the device itself rather than sending everything away to a far-off server, which also helps keep data more private.\n\n"
            "Data: Experts believe that within the next few years, there will be over a hundred billion of these smart gadgets working in homes all over the world.\n\n"
            "Conclusion: While these systems make life easier, they also require much better security to stop outsiders from interfering with things in the house."
        )),
        ("Network_Security", (
            "Introduction: With digital threats on the rise, many companies have moved to a security model that assumes every single person or device trying to connect is a potential risk until proven otherwise.\n\n"
            "Body 1: Old-fashioned firewalls aren't enough when hackers are so determined and skilled. Reliable defense now requires constantly watching every part of the network for any strange signs of activity.\n\n"
            "Body 2: While double-checking IDs and using complex codes is important, the human element is still the most likely way a hacker will get in, making regular staff training the most important part of the defense.\n\n"
            "Data: Recent studies show that the average cost for a business to fix the damage after a major hack is over four million dollars, a price most small companies can't afford.\n\n"
            "Conclusion: Keeping data safe isn't just about buying software; it's a never-ending job of staying one step ahead of the criminals."
        )),
        ("Remote_Compute", (
            "Introduction: Using shared remote servers has changed how startups and big companies build their websites by giving them instant access to massive power without having to buy any hardware.\n\n"
            "Body 1: Rented power - whether it's for simple file storage or complex math tasks - allows a business to only pay for what they use right now, rather than spending huge amounts of money on their own servers.\n\n"
            "Body 2: New ways of building websites that use 'containers' allow developers to update their apps many times a day without causing any crashes for the people visiting the site.\n\n"
            "Data: Reports show that over 90% of modern companies now use at least one of the big web-power providers for most of their daily work and customer services.\n\n"
            "Conclusion: Moving to these shared web systems is the only way for a modern business to stay fast and competitive in today's global market."
        ))
    ]
}

def create_pdf(folder, name, title, text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(5)
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 7, text)
    pdf.output(os.path.join(BASE_DIR, folder, f"{name}.pdf"))

def create_docx(folder, name, title, text):
    doc = Document()
    doc.add_heading(title, 0)
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(os.path.join(BASE_DIR, folder, f"{name}.docx"))

def create_txt(folder, name, text):
    with open(os.path.join(BASE_DIR, folder, f"{name}.txt"), "w", encoding='utf-8') as f:
        f.write(text)

# --- EXECUTION ---
print(f"Generating 7 sets of test files in {BASE_DIR}...")

for cat, items in SCENARIOS.items():
    print(f"  Generating category: {cat}")
    for i, (title, text) in enumerate(items):
        fname = f"{cat}_{i+1:02d}"
        create_pdf(cat, fname, title, text)
        create_docx(cat, fname, title, text)
        create_txt(cat, fname, text)

print("\n✅ All 14 sets (42 files) generated successfully!")
