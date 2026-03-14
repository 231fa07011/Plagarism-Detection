import os
import requests
import time
import urllib3

# Suppress insecure request warnings if verify=False is used
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

out_dir = r"c:\Users\MY PC\OneDrive\Desktop\Plagarism-Detection\sample_test_files\official_papers"
os.makedirs(out_dir, exist_ok=True)

downloads = [
    {"url": "https://aclanthology.org/N19-1423.pdf", "filename": "1_A4_Paper_BERT.pdf"},
    {"url": "https://aclanthology.org/P19-1452.pdf", "filename": "2_A4_Paper_RoBERTa.pdf"},
    {"url": "https://aclanthology.org/D14-1181.pdf", "filename": "3_A4_Paper_GloVe.pdf"},
    {"url": "https://aclanthology.org/P18-1017.pdf", "filename": "4_A4_Paper_NMT.pdf"},
    {"url": "https://aclanthology.org/2020.acl-main.442.pdf", "filename": "5_A4_Paper_BART.pdf"},
    {"url": "https://aclanthology.org/2020.emnlp-main.748.pdf", "filename": "6_A4_Paper_Longformer.pdf"},
    {"url": "https://aclanthology.org/D19-1410.pdf", "filename": "7_A4_Paper_SentenceBERT.pdf"},
    {"url": "https://aclanthology.org/W19-5004.pdf", "filename": "8_A4_Paper_BioBERT.pdf"},
    {"url": "https://aclanthology.org/N18-1101.pdf", "filename": "9_A4_Paper_ELMo.pdf"},
    {"url": "https://aclanthology.org/D15-1167.pdf", "filename": "10_A4_Paper_Doc2Vec.pdf"},
    {"url": "https://raw.githubusercontent.com/jbrownlee/Datasets/master/review_polarity/txt_sentoken/pos/cv000_29590.txt", "filename": "11_A5_Article_Review.txt"},
    {"url": "https://raw.githubusercontent.com/jbrownlee/Datasets/master/pima-indians-diabetes.names", "filename": "12_A5_Article_Pima.txt"},
    {"url": "https://raw.githubusercontent.com/jbrownlee/Datasets/master/housing.names", "filename": "13_A5_Article_Housing.txt"},
]

print(f"Starting downloads of {len(downloads) + 2} official papers and articles...")

success = 0
for i, item in enumerate(downloads, 1):
    output_path = os.path.join(out_dir, item["filename"])
    print(f"[{i}/15] Downloading {item['filename']}...")
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(item["url"], headers=headers, stream=True, timeout=20, verify=False)
        response.raise_for_status()
        
        with open(output_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        print("  -> Success!")
        success += 1
        time.sleep(1) # Prevent rate limiting
    except Exception as e:
        print(f"  -> Failed: {e}")

try:
    from docx import Document
    print("[14/15] Generating A3 formatted Official Document...")
    docx_path1 = os.path.join(out_dir, "14_A3_Conference_Paper_Mock.docx")
    doc = Document()
    doc.add_heading('Official Document Example: A3 Format', 0)
    doc.add_paragraph('This represents a generated official document targeting A3 layout structures. We use python-docx because large official A3 docx files are seldom available publicly without authenticating to academic portals.')
    doc.save(docx_path1)
    print("  -> Success!")
    success += 1

    print("[15/15] Generating A4 formatted Official Document...")
    docx_path2 = os.path.join(out_dir, "15_A4_Conference_Paper_Mock.docx")
    doc2 = Document()
    doc2.add_heading('Official Document Example: A4 Format', 0)
    doc2.add_paragraph('This represents a generated official document targeting A4 layout structures, analogous to standard thesis sizing.')
    doc2.save(docx_path2)
    print("  -> Success!")
    success += 1
except ImportError:
    print("python-docx not installed, skipping DOCX generation.")
except Exception as e:
    print(f"Failed to generate DOCX: {e}")

print(f"\nAll operations completed. {success}/15 files processed successfully.")
