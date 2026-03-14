import os
import requests
import urllib3
from docx import Document

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Get the directory where this script is located
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(BASE_DIR, "sample_test_files", "A3_A4_A5_Official_Papers")
os.makedirs(out_dir, exist_ok=True)

files = [
    ("https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf", "1_A3_Paper_W3C.pdf"),
    ("https://raw.githubusercontent.com/jbrownlee/Datasets/master/pima-indians-diabetes.names", "2_A5_Pima.txt"),
    ("https://raw.githubusercontent.com/jbrownlee/Datasets/master/housing.names", "3_A4_Housing.txt"),
    ("https://raw.githubusercontent.com/jbrownlee/Datasets/master/iris.names", "4_A3_Iris.txt"),
    ("https://www.rfc-editor.org/rfc/rfc2616.txt", "5_A4_HTTP11.txt"),
    ("https://www.rfc-editor.org/rfc/rfc8446.txt", "6_A5_TLS13.txt"),
    ("https://www.rfc-editor.org/rfc/rfc6749.txt", "7_A3_OAuth2.txt"),
    ("https://www.rfc-editor.org/rfc/rfc793.txt", "8_A4_TCP.txt"),
    ("https://www.rfc-editor.org/rfc/rfc1034.txt", "9_A5_DNS.txt"),
    ("https://www.rfc-editor.org/rfc/rfc1122.txt", "10_A3_HostRequirements.txt"),
]

print("--- DOWNLOADING FILES ---")
ok = 0
for i, (url, fname) in enumerate(files, 1):
    path = os.path.join(out_dir, fname)
    try:
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, stream=True, timeout=10, verify=False)
        r.raise_for_status()
        with open(path, "wb") as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
        size_kb = os.path.getsize(path) // 1024
        print(f"[{i}/15] SUCCESS -> {fname}  ({size_kb} KB)")
        ok += 1
    except Exception as e:
        print(f"[{i}/15] FAILED  -> {fname}  Error: {e}")

print()
print("--- GENERATING DOCX CONFERENCE PAPERS ---")
topics = [
    ("Neural Networks in Deep Learning", "A4"),
    ("Blockchain Security Protocols", "A3"),
    ("Quantum Computing Advances", "A5"),
    ("NLP and Transformer Models", "A4"),
    ("Autonomous Vehicles and AI", "A3"),
]

for j, (topic, fmt) in enumerate(topics, 11):
    safe_name = topic.replace(" ", "_")
    fname = f"{j}_{fmt}_{safe_name}.docx"
    path = os.path.join(out_dir, fname)
    doc = Document()
    doc.add_heading(f"{topic} - Official {fmt} Paper", 0)
    doc.add_heading("Abstract", 1)
    doc.add_paragraph(
        f"This paper presents a comprehensive study on {topic}. "
        f"Formatted in {fmt} paper size, this document conforms to standard academic submission guidelines."
    )
    doc.add_heading("1. Introduction", 1)
    doc.add_paragraph(
        f"Recent advances in {topic.lower()} have opened new research directions. "
        "This work surveys key milestones, methodologies, and open challenges."
    )
    doc.add_heading("2. Methodology", 1)
    doc.add_paragraph("Experimental evaluations were conducted on benchmark datasets using state-of-the-art models.")
    doc.add_heading("3. Conclusion", 1)
    doc.add_paragraph("This study highlights the growing importance of the field and outlines avenues for future work.")
    doc.save(path)
    size_kb = os.path.getsize(path) // 1024
    print(f"[{j}/15] SUCCESS -> {fname}  ({size_kb} KB)")
    ok += 1

print()
print(f"=== COMPLETED: {ok}/15 files ready ===")
print(f"Output folder: {out_dir}")
print()
print("Files in folder:")
for f in sorted(os.listdir(out_dir)):
    size = os.path.getsize(os.path.join(out_dir, f)) // 1024
    print(f"  {f}  ({size} KB)")
